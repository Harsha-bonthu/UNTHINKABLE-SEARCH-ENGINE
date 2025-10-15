# Continue creating all other files

files_content = {}

# 2. document_processor.py
files_content['document_processor.py'] = '''import fitz  # PyMuPDF
import docx
from pathlib import Path
from typing import List, Dict, Any
import asyncio
from langchain.text_splitters import RecursiveCharacterTextSplitter

class DocumentProcessor:
    """Handles document processing and text extraction for various file formats"""
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=["\\n\\n", "\\n", " ", ""]
        )
    
    async def process_document(self, file_path: str) -> List[Dict[str, Any]]:
        """Process a document and return chunks with metadata"""
        
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Extract text based on file type
        if file_path.suffix.lower() == '.pdf':
            text = await self._extract_text_from_pdf(str(file_path))
        elif file_path.suffix.lower() == '.docx':
            text = await self._extract_text_from_docx(str(file_path))
        elif file_path.suffix.lower() in ['.txt', '.md']:
            text = await self._extract_text_from_txt(str(file_path))
        else:
            raise ValueError(f"Unsupported file format: {file_path.suffix}")
        
        if not text.strip():
            raise ValueError("No text content found in document")
        
        # Split text into chunks
        chunks = self.text_splitter.split_text(text)
        
        # Create chunk objects with metadata
        chunk_objects = []
        for i, chunk in enumerate(chunks):
            if chunk.strip():  # Only include non-empty chunks
                chunk_objects.append({
                    'content': chunk.strip(),
                    'chunk_id': i,
                    'char_count': len(chunk),
                    'word_count': len(chunk.split()),
                })
        
        return chunk_objects
    
    async def _extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF using PyMuPDF"""
        try:
            doc = fitz.open(file_path)
            text = ""
            
            for page_num in range(doc.page_count):
                page = doc[page_num]
                text += page.get_text() + "\\n"
            
            doc.close()
            return text
            
        except Exception as e:
            raise Exception(f"Error extracting text from PDF: {str(e)}")
    
    async def _extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = docx.Document(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\\n"
            
            return text
            
        except Exception as e:
            raise Exception(f"Error extracting text from DOCX: {str(e)}")
    
    async def _extract_text_from_txt(self, file_path: str) -> str:
        """Extract text from TXT/MD file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
                
        except Exception as e:
            raise Exception(f"Error extracting text from TXT: {str(e)}")
    
    def get_document_stats(self, chunks: List[Dict[str, Any]]) -> Dict[str, Any]:
        """Get statistics about processed document"""
        if not chunks:
            return {}
        
        total_chars = sum(chunk['char_count'] for chunk in chunks)
        total_words = sum(chunk['word_count'] for chunk in chunks)
        
        return {
            'total_chunks': len(chunks),
            'total_characters': total_chars,
            'total_words': total_words,
            'avg_chunk_size': total_chars // len(chunks),
            'avg_words_per_chunk': total_words // len(chunks)
        }'''

# 3. vector_store.py
files_content['vector_store.py'] = '''import numpy as np
import faiss
import pickle
from pathlib import Path
from typing import List, Tuple, Dict, Any
import asyncio
from sentence_transformers import SentenceTransformer
import logging

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class VectorStore:
    """Vector database for storing and searching document embeddings"""
    
    def __init__(self, 
                 model_name: str = "all-MiniLM-L6-v2",
                 index_path: str = "vector_index.faiss",
                 metadata_path: str = "vector_metadata.pkl"):
        
        self.model_name = model_name
        self.index_path = index_path
        self.metadata_path = metadata_path
        self.dimension = 384  # all-MiniLM-L6-v2 embedding dimension
        
        # Initialize sentence transformer
        logger.info(f"Loading embedding model: {model_name}")
        self.encoder = SentenceTransformer(model_name)
        
        # Initialize FAISS index
        self.index = faiss.IndexFlatIP(self.dimension)  # Inner product for cosine similarity
        
        # Store metadata for each vector
        self.metadata = []
        self.texts = []
        
        # Load existing index if available
        self._load_index()
    
    async def add_documents(self, texts: List[str], metadatas: List[Dict[str, Any]]):
        """Add documents to the vector store"""
        
        if len(texts) != len(metadatas):
            raise ValueError("Number of texts and metadatas must match")
        
        if not texts:
            return
        
        logger.info(f"Adding {len(texts)} documents to vector store")
        
        # Generate embeddings
        embeddings = await self._generate_embeddings(texts)
        
        # Normalize embeddings for cosine similarity
        faiss.normalize_L2(embeddings)
        
        # Add to FAISS index
        self.index.add(embeddings)
        
        # Store texts and metadata
        self.texts.extend(texts)
        self.metadata.extend(metadatas)
        
        # Save updated index
        self._save_index()
        
        logger.info(f"Successfully added {len(texts)} documents. Total vectors: {self.index.ntotal}")
    
    async def similarity_search(self, 
                              query: str, 
                              k: int = 5) -> List[Tuple[str, Dict[str, Any], float]]:
        """Search for similar documents"""
        
        if self.index.ntotal == 0:
            return []
        
        # Generate query embedding
        query_embeddings = await self._generate_embeddings([query])
        query_embedding = query_embeddings[0:1]  # Keep as 2D array
        
        # Normalize query embedding
        faiss.normalize_L2(query_embedding)
        
        # Search
        k = min(k, self.index.ntotal)  # Don't search for more than available
        scores, indices = self.index.search(query_embedding, k)
        
        # Format results
        results = []
        for score, idx in zip(scores[0], indices[0]):
            if idx < len(self.texts) and idx != -1:  # Valid index
                results.append((
                    self.texts[idx],
                    self.metadata[idx],
                    float(score)
                ))
        
        return results
    
    async def _generate_embeddings(self, texts: List[str]) -> np.ndarray:
        """Generate embeddings for texts"""
        
        # Run in thread pool to avoid blocking
        loop = asyncio.get_event_loop()
        
        def encode_texts():
            return self.encoder.encode(
                texts, 
                convert_to_numpy=True,
                show_progress_bar=False
            )
        
        embeddings = await loop.run_in_executor(None, encode_texts)
        return embeddings.astype('float32')
    
    def _save_index(self):
        """Save FAISS index and metadata to disk"""
        try:
            # Save FAISS index
            faiss.write_index(self.index, self.index_path)
            
            # Save metadata and texts
            with open(self.metadata_path, 'wb') as f:
                pickle.dump({
                    'metadata': self.metadata,
                    'texts': self.texts,
                    'model_name': self.model_name
                }, f)
            
            logger.info(f"Saved vector index with {self.index.ntotal} vectors")
            
        except Exception as e:
            logger.error(f"Error saving index: {str(e)}")
    
    def _load_index(self):
        """Load FAISS index and metadata from disk"""
        try:
            index_path = Path(self.index_path)
            metadata_path = Path(self.metadata_path)
            
            if index_path.exists() and metadata_path.exists():
                # Load FAISS index
                self.index = faiss.read_index(self.index_path)
                
                # Load metadata and texts
                with open(self.metadata_path, 'rb') as f:
                    data = pickle.load(f)
                    self.metadata = data.get('metadata', [])
                    self.texts = data.get('texts', [])
                    saved_model = data.get('model_name', self.model_name)
                
                if saved_model != self.model_name:
                    logger.warning(f"Model mismatch: saved={saved_model}, current={self.model_name}")
                
                logger.info(f"Loaded vector index with {self.index.ntotal} vectors")
            else:
                logger.info("No existing index found, starting fresh")
                
        except Exception as e:
            logger.error(f"Error loading index: {str(e)}")
            logger.info("Starting with fresh index")
            self.index = faiss.IndexFlatIP(self.dimension)
            self.metadata = []
            self.texts = []
    
    def get_stats(self) -> Dict[str, Any]:
        """Get statistics about the vector store"""
        return {
            'total_vectors': int(self.index.ntotal),
            'dimension': self.dimension,
            'model_name': self.model_name,
            'total_texts': len(self.texts),
            'total_metadata': len(self.metadata)
        }
    
    def clear(self):
        """Clear all vectors and metadata"""
        self.index = faiss.IndexFlatIP(self.dimension)
        self.metadata = []
        self.texts = []
        
        # Remove saved files
        for path in [self.index_path, self.metadata_path]:
            if Path(path).exists():
                Path(path).unlink()
        
        logger.info("Cleared vector store")'''

# Create files
for filename, content in files_content.items():
    with open(filename, 'w') as f:
        f.write(content)
    print(f"✅ Created {filename}")

print("Created document_processor.py and vector_store.py")