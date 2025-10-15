import fitz  # PyMuPDF
import docx
from pathlib import Path
from typing import List, Dict, Any
import asyncio
from langchain.text_splitters import RecursiveCharacterTextSplitter

class DocumentProcessor:
    """Handles document processing and text extraction for various file formats"""

    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=["\n\n", "\n", " ", ""]
        )

    async def process_document(self, file_path: str) -> List[Dict[str, Any]]:
        """Process a document and return chunks with metadata"""

        file_path = Path(file_path)

        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        # Extract text based on file type
        if file_path.suffix.lower() == '.pdf':
            text = await self._extract_text_from_pdf(str(file_path))
        elif file_path.suffix.lower() == '.docx':
            text = await self._extract_text_from_docx(str(file_path))
        elif file_path.suffix.lower() in ['.txt', '.md']:
            text = await self._extract_text_from_txt(str(file_path))
        else:
            raise ValueError(f"Unsupported file format: {file_path.suffix}")

        if not text.strip():
            raise ValueError("No text content found in document")

        # Split text into chunks
        chunks = self.text_splitter.split_text(text)

        # Create chunk objects with metadata
        chunk_objects = []
        for i, chunk in enumerate(chunks):
            if chunk.strip():  # Only include non-empty chunks
                chunk_objects.append({
                    'content': chunk.strip(),
                    'chunk_id': i,
                    'char_count': len(chunk),
                    'word_count': len(chunk.split()),
                })

        return chunk_objects

    async def _extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF using PyMuPDF"""
        try:
            doc = fitz.open(file_path)
            text = ""

            for page_num in range(doc.page_count):
                page = doc[page_num]
                text += page.get_text() + "\n"

            doc.close()
            return text

        except Exception as e:
            raise Exception(f"Error extracting text from PDF: {str(e)}")

    async def _extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = docx.Document(file_path)
            text = ""

            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"

            return text

        except Exception as e:
            raise Exception(f"Error extracting text from DOCX: {str(e)}")

    async def _extract_text_from_txt(self, file_path: str) -> str:
        """Extract text from TXT/MD file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()

        except Exception as e:
            raise Exception(f"Error extracting text from TXT: {str(e)}")

    def get_document_stats(self, chunks: List[Dict[str, Any]]) -> Dict[str, Any]:
        """Get statistics about processed document"""
        if not chunks:
            return {}

        total_chars = sum(chunk['char_count'] for chunk in chunks)
        total_words = sum(chunk['word_count'] for chunk in chunks)

        return {
            'total_chunks': len(chunks),
            'total_characters': total_chars,
            'total_words': total_words,
            'avg_chunk_size': total_chars // len(chunks),
            'avg_words_per_chunk': total_words // len(chunks)
        }